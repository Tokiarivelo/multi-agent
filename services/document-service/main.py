import io
import os
import re
import csv
import json
from pathlib import Path
from typing import Optional, List, Any
from fastapi import FastAPI, HTTPException, UploadFile, File, Query
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn
import httpx

try:
    import matplotlib
    matplotlib.use("Agg")
    from matplotlib.figure import Figure
    from matplotlib.patches import Circle, FancyBboxPatch
    import numpy as np
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

try:
    import networkx as nx
    NETWORKX_AVAILABLE = True
except ImportError:
    NETWORKX_AVAILABLE = False

app = FastAPI(
    title="Document Generation Service",
    description=(
        "Generate, parse, and manage documents in multiple formats "
        "(PDF, DOCX, XLSX, MD, CSV, HTML, TXT, JSON). "
        "Also provides OCR for images and workspace file operations.\n\n"
        "**Swagger UI** is available at `/docs`; OpenAPI JSON at `/openapi.json`."
    ),
    version="1.0.0",
    contact={"name": "Multi-Agent Platform", "url": "https://github.com/multi-agent"},
    license_info={"name": "MIT"},
    openapi_tags=[
        {"name": "health", "description": "Liveness probe"},
        {"name": "documents", "description": "Generate and parse documents"},
        {"name": "files", "description": "Workspace file read / write / delete"},
    ],
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

SUPPORTED_FORMATS = ["pdf", "docx", "xlsx", "md", "csv", "html", "txt", "json"]

# Workspace root – all server-side path operations are restricted to this directory.
WORKSPACE_ROOT: str = os.environ.get("WORKSPACE_ROOT", os.getcwd())
FILE_SERVICE_URL: str = os.environ.get("FILE_SERVICE_URL", "http://localhost:3008")

MIME_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "md": "text/markdown",
    "csv": "text/csv",
    "html": "text/html",
    "txt": "text/plain",
    "json": "application/json",
}


class Section(BaseModel):
    heading: Optional[str] = None
    body: Optional[str] = None
    level: Optional[int] = 1


class TableData(BaseModel):
    headers: List[str] = []
    rows: List[List[Any]] = []


class DocumentMetadata(BaseModel):
    subject: Optional[str] = None
    keywords: Optional[str] = None
    company: Optional[str] = None


class ChartDataset(BaseModel):
    label: str = ""
    data: List[float] = []
    color: Optional[str] = None


class Chart(BaseModel):
    type: str = "bar"
    title: Optional[str] = None
    labels: List[str] = []
    datasets: List[ChartDataset] = []
    width: Optional[int] = 600
    height: Optional[int] = 400


class DiagramNode(BaseModel):
    id: str
    label: str
    shape: str = "box"   # box, circle, diamond, ellipse
    color: Optional[str] = None


class DiagramEdge(BaseModel):
    source: str
    target: str
    label: Optional[str] = None
    style: str = "solid"  # solid, dashed, dotted


class Diagram(BaseModel):
    type: str = "flowchart"  # flowchart, network, sequence, tree, mindmap
    title: Optional[str] = None
    nodes: List[DiagramNode] = []
    edges: List[DiagramEdge] = []
    direction: str = "TB"   # TB (top-bottom) or LR (left-right)
    width: Optional[int] = 700
    height: Optional[int] = 500


class GenerateRequest(BaseModel):
    format: str
    title: str = "Document"
    author: Optional[str] = None
    sections: Optional[List[Section]] = None
    table: Optional[TableData] = None
    charts: Optional[List[Chart]] = None
    diagrams: Optional[List[Diagram]] = None
    canvas_render: bool = False   # pre-render charts to PNG in HTML (offline / no CDN)
    metadata: Optional[DocumentMetadata] = None


def _render_chart_to_png(chart: Chart) -> bytes:
    if not MATPLOTLIB_AVAILABLE:
        raise HTTPException(status_code=500, detail="matplotlib not installed — charts unavailable")

    dpi = 100
    width_in = (chart.width or 600) / dpi
    height_in = (chart.height or 400) / dpi
    fig = Figure(figsize=(width_in, height_in), dpi=dpi)
    ax = fig.add_subplot(111)

    chart_type = chart.type.lower()
    labels = chart.labels
    datasets = chart.datasets

    if chart_type in ("pie", "doughnut") and datasets:
        ds = datasets[0]
        n = len(ds.data)
        pie_labels = labels[:n] if labels else [str(i) for i in range(n)]
        colors = [d.color for d in datasets if d.color] or None
        wedges, texts, autotexts = ax.pie(
            ds.data, labels=pie_labels, autopct="%1.1f%%",
            colors=colors if colors and len(colors) >= n else None,
        )
        if chart_type == "doughnut":
            ax.add_artist(Circle((0, 0), 0.5, fc="white"))
    elif chart_type in ("bar", "line", "scatter"):
        n_labels = len(labels) or max((len(d.data) for d in datasets), default=0)
        x = np.arange(n_labels)
        bar_width = 0.8 / max(len(datasets), 1)

        for i, ds in enumerate(datasets):
            y = ds.data
            color = ds.color or None
            lbl = ds.label or None
            if chart_type == "bar":
                offset = (i - (len(datasets) - 1) / 2.0) * bar_width
                ax.bar(x[:len(y)] + offset, y, width=bar_width, label=lbl, color=color)
            elif chart_type == "line":
                ax.plot(x[:len(y)], y, label=lbl, color=color, marker="o", markersize=4)
            else:
                ax.scatter(x[:len(y)], y, label=lbl, color=color, s=40)

        if labels:
            ax.set_xticks(x)
            ax.set_xticklabels(labels, rotation=30, ha="right", fontsize=8)
        if len(datasets) > 1 or (datasets and datasets[0].label):
            ax.legend(fontsize=8)
    else:
        for ds in datasets:
            ax.bar(range(len(ds.data)), ds.data, label=ds.label or None)
        if datasets and datasets[0].label:
            ax.legend(fontsize=8)

    if chart.title:
        ax.set_title(chart.title, fontsize=11)

    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    buf.seek(0)
    return buf.read()


def _chart_to_table_rows(chart: Chart) -> tuple[list[str], list[list]]:
    """Fallback: convert chart data to headers + rows for text-based formats."""
    headers = ["Label"] + [ds.label or f"Series {i+1}" for i, ds in enumerate(chart.datasets)]
    n = max((len(ds.data) for ds in chart.datasets), default=0)
    rows = []
    for i in range(n):
        label = chart.labels[i] if i < len(chart.labels) else str(i)
        row = [label] + [ds.data[i] if i < len(ds.data) else "" for ds in chart.datasets]
        rows.append(row)
    return headers, rows


# ── Diagram renderers ─────────────────────────────────────────────────────────

def _hierarchical_pos(G, direction: str = "TB") -> dict:
    """Level-based layout for DAGs without requiring graphviz binary."""
    try:
        topo = list(nx.topological_sort(G))
    except Exception:
        return nx.spring_layout(G, seed=42, k=2.5)

    levels: dict = {}
    for node in topo:
        preds = list(G.predecessors(node))
        levels[node] = 0 if not preds else max(levels.get(p, 0) for p in preds) + 1

    level_groups: dict = {}
    for node, lvl in levels.items():
        level_groups.setdefault(lvl, []).append(node)

    pos: dict = {}
    for lvl, nodes in level_groups.items():
        n = len(nodes)
        for i, node in enumerate(nodes):
            x = (i - (n - 1) / 2.0) * 2.0
            y = -float(lvl) * 2.0
            pos[node] = (y, x) if direction == "LR" else (x, y)
    return pos


def _render_sequence_diagram_png(diagram: Diagram, w: float, h: float, dpi: int) -> bytes:
    fig = Figure(figsize=(w, h), dpi=dpi)
    ax = fig.add_subplot(111)
    ax.axis("off")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)

    actors = [n.id for n in diagram.nodes]
    if not actors:
        all_ids = {e.source for e in diagram.edges} | {e.target for e in diagram.edges}
        actors = sorted(all_ids)
    actor_label = {n.id: n.label for n in diagram.nodes}

    n_actors = len(actors)
    if n_actors == 0:
        buf = io.BytesIO()
        fig.savefig(buf, format="png")
        buf.seek(0)
        return buf.read()

    actor_x = {a: (i + 0.5) / n_actors for i, a in enumerate(actors)}

    for aid, x in actor_x.items():
        lbl = actor_label.get(aid, aid)
        ax.text(x, 0.95, lbl, ha="center", va="center", fontsize=9,
                bbox=dict(boxstyle="round,pad=0.3", facecolor="#4f46e5", edgecolor="none"),
                color="white", fontweight="bold")
        ax.plot([x, x], [0.05, 0.91], color="#cbd5e1", linewidth=1, linestyle="dashed")

    n_msgs = len(diagram.edges)
    y_start, y_end = 0.85, 0.10
    y_step = (y_start - y_end) / max(n_msgs, 1)
    for i, edge in enumerate(diagram.edges):
        y = y_start - i * y_step
        x1 = actor_x.get(edge.source, 0.1)
        x2 = actor_x.get(edge.target, 0.9)
        ls = "dashed" if edge.style == "dashed" else "solid"
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="->", color="#334155", lw=1.5, linestyle=ls))
        if edge.label:
            ax.text((x1 + x2) / 2, y + 0.025, edge.label,
                    ha="center", va="bottom", fontsize=8, color="#475569")

    if diagram.title:
        ax.set_title(diagram.title, fontsize=11)

    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    buf.seek(0)
    return buf.read()


_NODE_COLORS = ["#4f46e5", "#0ea5e9", "#10b981", "#f59e0b", "#ef4444",
                "#8b5cf6", "#06b6d4", "#84cc16", "#f97316", "#ec4899"]


def _render_diagram_to_png(diagram: Diagram) -> bytes:
    if not MATPLOTLIB_AVAILABLE:
        raise HTTPException(status_code=500, detail="matplotlib not installed — diagrams unavailable")

    dpi = 100
    w = (diagram.width or 700) / dpi
    h = (diagram.height or 500) / dpi

    if diagram.type == "sequence":
        return _render_sequence_diagram_png(diagram, w, h, dpi)

    if not NETWORKX_AVAILABLE:
        raise HTTPException(status_code=500, detail="networkx not installed — graph diagrams unavailable")

    fig = Figure(figsize=(w, h), dpi=dpi)
    ax = fig.add_subplot(111)
    ax.axis("off")
    ax.set_facecolor("white")

    is_directed = diagram.type in ("flowchart", "tree")
    G = nx.DiGraph() if is_directed else nx.Graph()

    node_label_map: dict = {}
    node_color_map: dict = {}
    node_shape_map: dict = {}
    for i, node in enumerate(diagram.nodes):
        G.add_node(node.id)
        node_label_map[node.id] = node.label
        node_color_map[node.id] = node.color or _NODE_COLORS[i % len(_NODE_COLORS)]
        node_shape_map[node.id] = node.shape or "box"

    edge_label_map: dict = {}
    edge_style_map: dict = {}
    for edge in diagram.edges:
        if edge.source not in G:
            G.add_node(edge.source)
            node_label_map[edge.source] = edge.source
            node_color_map[edge.source] = _NODE_COLORS[len(G) % len(_NODE_COLORS)]
        if edge.target not in G:
            G.add_node(edge.target)
            node_label_map[edge.target] = edge.target
            node_color_map[edge.target] = _NODE_COLORS[len(G) % len(_NODE_COLORS)]
        G.add_edge(edge.source, edge.target)
        if edge.label:
            edge_label_map[(edge.source, edge.target)] = edge.label
        edge_style_map[(edge.source, edge.target)] = edge.style

    if len(G.nodes) == 0:
        buf = io.BytesIO()
        fig.savefig(buf, format="png")
        buf.seek(0)
        return buf.read()

    # Choose layout
    if diagram.type in ("tree", "flowchart") and is_directed:
        try:
            pos = _hierarchical_pos(G, diagram.direction)
        except Exception:
            pos = nx.spring_layout(G, seed=42, k=2.5)
    elif diagram.type == "mindmap":
        pos = nx.shell_layout(G)
    else:
        pos = nx.spring_layout(G, seed=42, k=2.5)

    node_ids = list(G.nodes())
    colors = [node_color_map.get(n, "#4f46e5") for n in node_ids]
    labels = {n: node_label_map.get(n, n) for n in node_ids}

    # Separate diamond (decision) nodes — drawn manually
    diamond_nodes = [n for n in node_ids if node_shape_map.get(n) == "diamond"]
    circle_nodes = [n for n in node_ids if node_shape_map.get(n) in ("circle", "ellipse")]
    box_nodes = [n for n in node_ids if n not in diamond_nodes and n not in circle_nodes]

    nx.draw_networkx_nodes(G, pos, nodelist=box_nodes, ax=ax,
                           node_color=[node_color_map.get(n, "#4f46e5") for n in box_nodes],
                           node_size=2200, node_shape="s", alpha=0.92)
    nx.draw_networkx_nodes(G, pos, nodelist=circle_nodes, ax=ax,
                           node_color=[node_color_map.get(n, "#0ea5e9") for n in circle_nodes],
                           node_size=2200, node_shape="o", alpha=0.92)
    nx.draw_networkx_nodes(G, pos, nodelist=diamond_nodes, ax=ax,
                           node_color=[node_color_map.get(n, "#f59e0b") for n in diamond_nodes],
                           node_size=2200, node_shape="D", alpha=0.92)

    nx.draw_networkx_labels(G, pos, labels=labels, ax=ax,
                            font_color="white", font_size=8, font_weight="bold")

    solid_edges = [(u, v) for u, v in G.edges() if edge_style_map.get((u, v), "solid") != "dashed"]
    dashed_edges = [(u, v) for u, v in G.edges() if edge_style_map.get((u, v)) == "dashed"]

    draw_kw = dict(ax=ax, arrows=is_directed, arrowsize=16,
                   edge_color="#64748b", width=1.8, connectionstyle="arc3,rad=0.05")
    if solid_edges:
        nx.draw_networkx_edges(G, pos, edgelist=solid_edges, **draw_kw)
    if dashed_edges:
        nx.draw_networkx_edges(G, pos, edgelist=dashed_edges, style="dashed", **draw_kw)
    if edge_label_map:
        nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_label_map, ax=ax, font_size=8)

    if diagram.title:
        ax.set_title(diagram.title, fontsize=11, pad=10)

    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    buf.seek(0)
    return buf.read()


def _diagram_to_text_lines(diagram: Diagram) -> list[str]:
    heading = diagram.title or f"{diagram.type.capitalize()} Diagram"
    lines = [heading, "-" * len(heading)]
    lines.append(f"Type: {diagram.type}  Direction: {diagram.direction}")
    if diagram.nodes:
        lines.append("Nodes: " + ", ".join(f"{n.id}({n.label})" for n in diagram.nodes))
    for edge in diagram.edges:
        arrow = "-->" if edge.style != "dashed" else "- ->"
        lbl = f" [{edge.label}]" if edge.label else ""
        lines.append(f"  {edge.source} {arrow} {edge.target}{lbl}")
    return lines


def _generate_pdf(req: GenerateRequest) -> bytes:
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm)
        styles = getSampleStyleSheet()
        story = []

        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Title"],
            fontSize=22,
            spaceAfter=20,
        )
        h1_style = ParagraphStyle(
            "CustomH1",
            parent=styles["Heading1"],
            fontSize=14,
            spaceAfter=8,
            spaceBefore=16,
        )
        body_style = styles["BodyText"]
        body_style.leading = 16

        story.append(Paragraph(req.title, title_style))
        if req.author:
            story.append(Paragraph(f"Author: {req.author}", styles["Normal"]))
        story.append(Spacer(1, 0.5 * cm))

        for section in req.sections or []:
            if section.heading:
                story.append(Paragraph(section.heading, h1_style))
            if section.body:
                story.append(Paragraph(section.body, body_style))
            story.append(Spacer(1, 0.3 * cm))

        if req.table and req.table.headers:
            table_data = [req.table.headers] + [list(map(str, r)) for r in req.table.rows]
            t = Table(table_data, repeatRows=1)
            t.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4f46e5")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 10),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(t)

        from reportlab.platypus import Image as RLImage

        for chart in req.charts or []:
            story.append(Spacer(1, 0.5 * cm))
            if chart.title:
                story.append(Paragraph(chart.title, h1_style))
            png_bytes = _render_chart_to_png(chart)
            img = RLImage(io.BytesIO(png_bytes), width=14 * cm,
                          height=(14 * cm * (chart.height or 400)) / (chart.width or 600))
            story.append(img)

        for diagram in req.diagrams or []:
            story.append(Spacer(1, 0.5 * cm))
            if diagram.title:
                story.append(Paragraph(diagram.title, h1_style))
            png_bytes = _render_diagram_to_png(diagram)
            img = RLImage(io.BytesIO(png_bytes), width=14 * cm,
                          height=(14 * cm * (diagram.height or 500)) / (diagram.width or 700))
            story.append(img)

        doc.build(story)
        return buf.getvalue()
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"reportlab not installed: {e}")


def _generate_docx(req: GenerateRequest) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        core = doc.core_properties
        core.title = req.title
        if req.author:
            core.author = req.author
        if req.metadata:
            if req.metadata.subject:
                core.subject = req.metadata.subject
            if req.metadata.keywords:
                core.keywords = req.metadata.keywords
            if req.metadata.company:
                core.company = req.metadata.company

        title_para = doc.add_heading(req.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if req.author:
            p = doc.add_paragraph(f"Author: {req.author}")
            p.runs[0].italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        for section in req.sections or []:
            if section.heading:
                doc.add_heading(section.heading, level=section.level or 1)
            if section.body:
                doc.add_paragraph(section.body)

        if req.table and req.table.headers:
            doc.add_paragraph()
            t = doc.add_table(rows=1, cols=len(req.table.headers))
            t.style = "Table Grid"
            hdr = t.rows[0].cells
            for i, h in enumerate(req.table.headers):
                hdr[i].text = str(h)
                run = hdr[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(79, 70, 229)
            for row in req.table.rows:
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)

        from docx.shared import Inches

        for chart in req.charts or []:
            doc.add_paragraph()
            if chart.title:
                doc.add_heading(chart.title, level=2)
            png_bytes = _render_chart_to_png(chart)
            doc.add_picture(io.BytesIO(png_bytes), width=Inches(5))

        for diagram in req.diagrams or []:
            doc.add_paragraph()
            if diagram.title:
                doc.add_heading(diagram.title, level=2)
            png_bytes = _render_diagram_to_png(diagram)
            doc.add_picture(io.BytesIO(png_bytes), width=Inches(5.5))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"python-docx not installed: {e}")


def _generate_xlsx(req: GenerateRequest) -> bytes:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = req.title[:31]

        wb.properties.title = req.title
        if req.author:
            wb.properties.creator = req.author

        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill("solid", fgColor="4F46E5")
        header_align = Alignment(horizontal="center", vertical="center")
        thin = Side(style="thin", color="E2E8F0")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        row_offset = 1
        if req.sections:
            for section in req.sections:
                if section.heading:
                    cell = ws.cell(row=row_offset, column=1, value=section.heading)
                    cell.font = Font(bold=True, size=13)
                    row_offset += 1
                if section.body:
                    cell = ws.cell(row=row_offset, column=1, value=section.body)
                    row_offset += 1
                row_offset += 1

        if req.table and req.table.headers:
            for col_idx, header in enumerate(req.table.headers, start=1):
                cell = ws.cell(row=row_offset, column=col_idx, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_align
                cell.border = border
            row_offset += 1

            alt_fill = PatternFill("solid", fgColor="F1F5F9")
            for r_idx, row in enumerate(req.table.rows):
                fill = alt_fill if r_idx % 2 == 1 else None
                for col_idx, val in enumerate(row, start=1):
                    cell = ws.cell(row=row_offset, column=col_idx, value=val)
                    cell.border = border
                    if fill:
                        cell.fill = fill
                row_offset += 1

            for col_idx in range(1, len(req.table.headers) + 1):
                col_letter = get_column_letter(col_idx)
                ws.column_dimensions[col_letter].width = 18

        for chart in req.charts or []:
            from openpyxl.chart import BarChart, LineChart, PieChart, Reference
            from openpyxl.chart.series import DataPoint

            # Write chart data starting 2 rows after current content
            data_start_row = row_offset + 2
            ws.cell(row=data_start_row, column=1, value="Label")
            for ds_idx, ds in enumerate(chart.datasets):
                ws.cell(row=data_start_row, column=ds_idx + 2, value=ds.label or f"Series {ds_idx + 1}")

            n_points = max((len(ds.data) for ds in chart.datasets), default=0)
            for i in range(n_points):
                label = chart.labels[i] if i < len(chart.labels) else str(i + 1)
                ws.cell(row=data_start_row + 1 + i, column=1, value=label)
                for ds_idx, ds in enumerate(chart.datasets):
                    val = ds.data[i] if i < len(ds.data) else None
                    ws.cell(row=data_start_row + 1 + i, column=ds_idx + 2, value=val)

            chart_type = chart.type.lower()
            if chart_type == "line":
                xl_chart = LineChart()
            elif chart_type == "pie":
                xl_chart = PieChart()
            else:
                xl_chart = BarChart()
                xl_chart.type = "col"
                xl_chart.grouping = "clustered"

            xl_chart.title = chart.title or ""
            xl_chart.style = 10

            data_ref = Reference(
                ws,
                min_col=2,
                max_col=1 + len(chart.datasets),
                min_row=data_start_row,
                max_row=data_start_row + n_points,
            )
            cats_ref = Reference(
                ws,
                min_col=1,
                min_row=data_start_row + 1,
                max_row=data_start_row + n_points,
            )
            xl_chart.add_data(data_ref, titles_from_data=True)
            if chart_type != "pie":
                xl_chart.set_categories(cats_ref)

            chart_anchor = f"A{data_start_row + n_points + 3}"
            ws.add_chart(xl_chart, chart_anchor)
            row_offset = data_start_row + n_points + 20

        for diagram in req.diagrams or []:
            from openpyxl.drawing.image import Image as XLImage
            import tempfile, os as _os

            png_bytes = _render_diagram_to_png(diagram)
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(png_bytes)
                tmp_path = tmp.name
            try:
                xl_img = XLImage(tmp_path)
                xl_img.width = 480
                xl_img.height = int(480 * (diagram.height or 500) / (diagram.width or 700))
                ws.add_image(xl_img, f"A{row_offset + 2}")
                row_offset += int(xl_img.height / 20) + 4
            finally:
                _os.unlink(tmp_path)

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"openpyxl not installed: {e}")


def _generate_md(req: GenerateRequest) -> bytes:
    lines = [f"# {req.title}", ""]
    if req.author:
        lines += [f"**Author:** {req.author}", ""]
    if req.metadata:
        if req.metadata.subject:
            lines += [f"**Subject:** {req.metadata.subject}", ""]
    lines.append("---")
    lines.append("")
    for section in req.sections or []:
        if section.heading:
            prefix = "#" * (section.level or 1 + 1)
            lines.append(f"{prefix} {section.heading}")
            lines.append("")
        if section.body:
            lines.append(section.body)
            lines.append("")
    if req.table and req.table.headers:
        lines.append("| " + " | ".join(req.table.headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(req.table.headers)) + " |")
        for row in req.table.rows:
            lines.append("| " + " | ".join(str(v) for v in row) + " |")
        lines.append("")
    for chart in req.charts or []:
        heading = chart.title or f"{chart.type.capitalize()} Chart"
        lines += [f"## {heading}", ""]
        headers, rows = _chart_to_table_rows(chart)
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows:
            lines.append("| " + " | ".join(str(v) for v in row) + " |")
        lines.append("")
    for diagram in req.diagrams or []:
        lines.append("")
        for line in _diagram_to_text_lines(diagram):
            lines.append(line)
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def _generate_csv(req: GenerateRequest) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    if req.table:
        if req.table.headers:
            writer.writerow(req.table.headers)
        for row in req.table.rows:
            writer.writerow(row)
    elif req.sections:
        writer.writerow(["Heading", "Body"])
        for s in req.sections:
            writer.writerow([s.heading or "", s.body or ""])
    for chart in req.charts or []:
        writer.writerow([])
        writer.writerow([f"# {chart.title or chart.type} chart"])
        headers, rows = _chart_to_table_rows(chart)
        writer.writerow(headers)
        for row in rows:
            writer.writerow(row)
    for diagram in req.diagrams or []:
        writer.writerow([])
        writer.writerow([f"# {diagram.title or diagram.type} diagram"])
        for line in _diagram_to_text_lines(diagram):
            writer.writerow([line])
    return buf.getvalue().encode("utf-8")


def _generate_html(req: GenerateRequest) -> bytes:
    sections_html = ""
    for section in req.sections or []:
        if section.heading:
            level = section.level or 2
            sections_html += f"<h{level}>{section.heading}</h{level}>\n"
        if section.body:
            for para in section.body.split("\n\n"):
                sections_html += f"<p>{para}</p>\n"

    table_html = ""
    if req.table and req.table.headers:
        headers = "".join(f"<th>{h}</th>" for h in req.table.headers)
        rows_html = ""
        for row in req.table.rows:
            cells = "".join(f"<td>{v}</td>" for v in row)
            rows_html += f"<tr>{cells}</tr>\n"
        table_html = f"""
        <table>
          <thead><tr>{headers}</tr></thead>
          <tbody>{rows_html}</tbody>
        </table>"""

    import base64 as _b64

    charts_html = ""
    charts_scripts = ""

    for idx, chart in enumerate(req.charts or []):
        title_tag = f"<h2>{chart.title}</h2>" if chart.title else ""
        if req.canvas_render:
            # Server-side canvas render — embed pre-rendered PNG as <img>
            png_bytes = _render_chart_to_png(chart)
            b64 = _b64.b64encode(png_bytes).decode()
            charts_html += f"""
  <div class="media-container">
    {title_tag}
    <img src="data:image/png;base64,{b64}" width="{chart.width or 600}" height="{chart.height or 400}" alt="{chart.title or 'chart'}" style="max-width:100%;border-radius:8px;">
  </div>"""
        else:
            canvas_id = f"chart_{idx}"
            chart_type_js = "doughnut" if chart.type == "doughnut" else chart.type.lower()
            datasets_js = json.dumps([
                {
                    "label": ds.label,
                    "data": ds.data,
                    "backgroundColor": ds.color or f"hsl({idx * 60 + i * 37}, 65%, 55%)",
                    "borderColor": ds.color or f"hsl({idx * 60 + i * 37}, 65%, 45%)",
                    "fill": False,
                    "tension": 0.3,
                }
                for i, ds in enumerate(chart.datasets)
            ])
            labels_js = json.dumps(chart.labels)
            title_js = json.dumps(chart.title or "")
            charts_html += f"""
  <div class="media-container">
    {title_tag}
    <canvas id="{canvas_id}" width="{chart.width or 600}" height="{chart.height or 400}"></canvas>
  </div>"""
            charts_scripts += f"""
  new Chart(document.getElementById({json.dumps(canvas_id)}), {{
    type: {json.dumps(chart_type_js)},
    data: {{ labels: {labels_js}, datasets: {datasets_js} }},
    options: {{
      responsive: false,
      plugins: {{ legend: {{ position: "bottom" }}, title: {{ display: !!{title_js}, text: {title_js} }} }},
    }},
  }});"""

    diagrams_html = ""
    for diagram in req.diagrams or []:
        png_bytes = _render_diagram_to_png(diagram)
        b64 = _b64.b64encode(png_bytes).decode()
        title_tag = f"<h2>{diagram.title}</h2>" if diagram.title else ""
        diagrams_html += f"""
  <div class="media-container">
    {title_tag}
    <img src="data:image/png;base64,{b64}" width="{diagram.width or 700}" height="{diagram.height or 500}" alt="{diagram.title or 'diagram'}" style="max-width:100%;border-radius:8px;">
  </div>"""

    chartjs_script = (
        '<script src="https://cdn.jsdelivr.net/npm/chart.js@4/dist/chart.umd.min.js"></script>'
        if req.charts and not req.canvas_render
        else ""
    )

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{req.title}</title>
  {chartjs_script}
  <style>
    body {{ font-family: system-ui, sans-serif; max-width: 900px; margin: 2rem auto; padding: 0 1rem; color: #1e293b; }}
    h1 {{ color: #4f46e5; border-bottom: 2px solid #e2e8f0; padding-bottom: .5rem; }}
    h2, h3 {{ color: #334155; }}
    p {{ line-height: 1.7; }}
    table {{ width: 100%; border-collapse: collapse; margin-top: 1rem; }}
    th {{ background: #4f46e5; color: #fff; padding: .6rem 1rem; text-align: left; }}
    td {{ padding: .5rem 1rem; border-bottom: 1px solid #e2e8f0; }}
    tr:nth-child(even) td {{ background: #f8fafc; }}
    .meta {{ color: #64748b; font-size: .9rem; margin-bottom: 1.5rem; }}
    .media-container {{ margin: 2rem 0; }}
  </style>
</head>
<body>
  <h1>{req.title}</h1>
  <p class="meta">{f'Author: {req.author}' if req.author else ''}</p>
  {sections_html}
  {table_html}
  {charts_html}
  {diagrams_html}
  {f'<script>{charts_scripts}</script>' if charts_scripts else ''}
</body>
</html>"""
    return html.encode("utf-8")


def _generate_txt(req: GenerateRequest) -> bytes:
    lines = [req.title, "=" * len(req.title), ""]
    if req.author:
        lines += [f"Author: {req.author}", ""]
    for section in req.sections or []:
        if section.heading:
            lines += [section.heading, "-" * len(section.heading)]
        if section.body:
            lines.append(section.body)
        lines.append("")
    if req.table and req.table.headers:
        col_widths = [max(len(str(h)), *(len(str(r[i])) for r in req.table.rows), 0) + 2 for i, h in enumerate(req.table.headers)]
        header_row = " | ".join(str(h).ljust(col_widths[i]) for i, h in enumerate(req.table.headers))
        lines.append(header_row)
        lines.append("-" * len(header_row))
        for row in req.table.rows:
            lines.append(" | ".join(str(v).ljust(col_widths[i]) for i, v in enumerate(row)))
    for chart in req.charts or []:
        heading = chart.title or f"{chart.type.capitalize()} Chart"
        lines += ["", heading, "-" * len(heading)]
        headers, rows = _chart_to_table_rows(chart)
        col_widths = [max(len(str(h)), *(len(str(r[i])) for r in rows), 0) + 2 for i, h in enumerate(headers)]
        header_row = " | ".join(str(h).ljust(col_widths[i]) for i, h in enumerate(headers))
        lines.append(header_row)
        lines.append("-" * len(header_row))
        for row in rows:
            lines.append(" | ".join(str(v).ljust(col_widths[i]) for i, v in enumerate(row)))
    for diagram in req.diagrams or []:
        lines.append("")
        lines.extend(_diagram_to_text_lines(diagram))
    return "\n".join(lines).encode("utf-8")


def _generate_json_doc(req: GenerateRequest) -> bytes:
    data = {
        "title": req.title,
        "author": req.author,
        "sections": [s.model_dump() for s in (req.sections or [])],
        "table": req.table.model_dump() if req.table else None,
        "charts": [c.model_dump() for c in (req.charts or [])],
        "diagrams": [d.model_dump() for d in (req.diagrams or [])],
        "metadata": req.metadata.model_dump() if req.metadata else None,
    }
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")


GENERATORS = {
    "pdf": _generate_pdf,
    "docx": _generate_docx,
    "xlsx": _generate_xlsx,
    "md": _generate_md,
    "csv": _generate_csv,
    "html": _generate_html,
    "txt": _generate_txt,
    "json": _generate_json_doc,
}


@app.get("/api/documents/health", tags=["health"], summary="Health check")
def health():
    return {"status": "ok", "service": "document-service"}


@app.get("/api/documents/formats", tags=["documents"], summary="List supported output formats")
def list_formats():
    return {
        "formats": [
            {"id": "pdf", "label": "PDF", "mime": MIME_TYPES["pdf"], "ext": ".pdf"},
            {"id": "docx", "label": "Word (DOCX)", "mime": MIME_TYPES["docx"], "ext": ".docx"},
            {"id": "xlsx", "label": "Excel (XLSX)", "mime": MIME_TYPES["xlsx"], "ext": ".xlsx"},
            {"id": "md", "label": "Markdown", "mime": MIME_TYPES["md"], "ext": ".md"},
            {"id": "csv", "label": "CSV", "mime": MIME_TYPES["csv"], "ext": ".csv"},
            {"id": "html", "label": "HTML", "mime": MIME_TYPES["html"], "ext": ".html"},
            {"id": "txt", "label": "Plain Text", "mime": MIME_TYPES["txt"], "ext": ".txt"},
            {"id": "json", "label": "JSON", "mime": MIME_TYPES["json"], "ext": ".json"},
        ]
    }


@app.post(
    "/api/documents/generate",
    tags=["documents"],
    summary="Generate a document",
    description="Generate a document in the requested format. Returns the file as a binary stream with an appropriate Content-Disposition header.",
    response_class=StreamingResponse,
)
def generate_document(req: GenerateRequest):
    fmt = req.format.lower()
    if fmt not in SUPPORTED_FORMATS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported format '{fmt}'. Supported: {', '.join(SUPPORTED_FORMATS)}",
        )

    generator = GENERATORS[fmt]
    content = generator(req)

    safe_title = "".join(c if c.isalnum() or c in "._- " else "_" for c in req.title).strip()
    filename = f"{safe_title or 'document'}.{fmt}"

    return StreamingResponse(
        io.BytesIO(content),
        media_type=MIME_TYPES[fmt],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── Parse models ────────────────────────────────────────────────────────────

class ParsePathRequest(BaseModel):
    path: str
    encoding: Optional[str] = "utf-8"


class ParseImagePathRequest(BaseModel):
    path: str


class WriteFileRequest(BaseModel):
    path: str
    content: str
    encoding: Optional[str] = "utf-8"
    userId: Optional[str] = None
    workspaceRoot: Optional[str] = None


class DeleteFileRequest(BaseModel):
    path: str


# ─── Parse helpers ────────────────────────────────────────────────────────────

def _safe_workspace_path(file_path: str, check_exists: bool = True) -> str:
    """Resolve *file_path* inside WORKSPACE_ROOT and reject escapes via symlinks / '..'."""
    path, is_safe = _resolve_path(file_path)
    if not is_safe:
        raise HTTPException(
            status_code=403,
            detail="Access denied: path is outside the allowed workspace root",
        )
    if check_exists and not os.path.exists(path):
        raise HTTPException(status_code=404, detail=f"File not found: {file_path}")
    return path


def _resolve_path(file_path: str, workspace_root: Optional[str] = None) -> tuple[str, bool]:
    """Resolve *file_path* against *workspace_root* (or WORKSPACE_ROOT) and check containment."""
    root = workspace_root or WORKSPACE_ROOT
    if os.path.isabs(file_path):
        candidate = file_path
    else:
        candidate = os.path.join(root, file_path)

    real_root = os.path.realpath(root)
    real_candidate = os.path.realpath(candidate)

    is_safe = real_candidate.startswith(real_root + os.sep) or real_candidate == real_root
    return real_candidate, is_safe


async def _upload_to_cloud(
    filename: str,
    content: bytes,
    user_id: str,
    mime_type: str = "text/plain",
) -> dict:
    """Upload content to cloud storage using presigned PUT URL (best practice).

    Two-step flow:
    1. POST /api/files/initiate-upload → receives presigned PUT URL + file record.
    2. PUT bytes directly to MinIO via presigned URL (no buffering through file-service).
    """
    async with httpx.AsyncClient() as http:
        # Step 1: initiate upload — get presigned PUT URL
        init_resp = await http.post(
            f"{FILE_SERVICE_URL}/api/files/initiate-upload",
            params={"userId": user_id},
            json={"originalName": filename, "mimeType": mime_type, "size": len(content)},
            timeout=15.0,
        )
        if not init_resp.is_success:
            raise HTTPException(
                status_code=init_resp.status_code,
                detail=f"Cloud upload initiation failed: {init_resp.text}",
            )
        data = init_resp.json()
        upload_url: str = data["uploadUrl"]
        record: dict = data["record"]

        # Step 2: PUT bytes directly to MinIO presigned URL (bypasses file-service buffer)
        put_resp = await http.put(
            upload_url,
            content=content,
            headers={"Content-Type": mime_type},
            timeout=60.0,
        )
        if not put_resp.is_success:
            raise HTTPException(
                status_code=put_resp.status_code,
                detail=f"Cloud upload PUT failed: {put_resp.text}",
            )

    return {"url": record["url"], "fileId": record.get("id", ""), "storedAs": record.get("storedName", filename)}


def _detect_ext(filename: str) -> str:
    return Path(filename).suffix.lower().lstrip(".")


def _parse_pdf_bytes(content: bytes) -> dict:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return {"type": "pdf", "text": "\n\n".join(pages), "pages": len(pages)}
    except ImportError:
        raise HTTPException(status_code=500, detail="pdfplumber not installed")


def _parse_docx_bytes(content: bytes) -> dict:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = [
            [[cell.text for cell in row.cells] for row in table.rows]
            for table in doc.tables
        ]
        return {"type": "docx", "text": "\n".join(paragraphs), "paragraphs": paragraphs, "tables": tables}
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")


def _parse_xlsx_bytes(content: bytes) -> dict:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(content), data_only=True)
        sheets: dict = {}
        for name in wb.sheetnames:
            ws = wb[name]
            sheets[name] = [
                [str(cell.value) if cell.value is not None else "" for cell in row]
                for row in ws.iter_rows()
            ]
        return {"type": "xlsx", "sheets": sheets}
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl not installed")


def _parse_csv_bytes(content: bytes, encoding: str) -> dict:
    text = content.decode(encoding, errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    headers = rows[0] if rows else []
    data = rows[1:] if len(rows) > 1 else []
    return {"type": "csv", "headers": headers, "rows": data, "row_count": len(data)}


def _parse_html_bytes(content: bytes, encoding: str) -> dict:
    text = content.decode(encoding, errors="replace")
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(text, "html.parser")
        plain = soup.get_text(separator="\n", strip=True)
        title_tag = soup.find("title")
        return {"type": "html", "text": plain, "title": title_tag.get_text() if title_tag else None}
    except ImportError:
        plain = re.sub(r"<[^>]+>", " ", text)
        plain = re.sub(r"\s+", " ", plain).strip()
        return {"type": "html", "text": plain, "title": None}


def _parse_file_bytes(filename: str, content: bytes, encoding: str = "utf-8") -> dict:
    ext = _detect_ext(filename)
    if ext == "pdf":
        return _parse_pdf_bytes(content)
    if ext == "docx":
        return _parse_docx_bytes(content)
    if ext in ("xlsx", "xls"):
        return _parse_xlsx_bytes(content)
    if ext == "csv":
        return _parse_csv_bytes(content, encoding)
    if ext == "json":
        raw = content.decode(encoding, errors="replace")
        return {"type": "json", "data": json.loads(raw), "raw": raw}
    if ext in ("html", "htm"):
        return _parse_html_bytes(content, encoding)
    # txt, md, and everything else
    text = content.decode(encoding, errors="replace")
    return {"type": ext or "text", "text": text, "length": len(text)}


def _parse_image_bytes(content: bytes, filename: str) -> dict:
    try:
        from PIL import Image as PilImage
        img = PilImage.open(io.BytesIO(content))
        result: dict = {
            "type": "image",
            "format": img.format,
            "mode": img.mode,
            "width": img.width,
            "height": img.height,
            "filename": filename,
        }
        try:
            import pytesseract
            result["ocr_text"] = pytesseract.image_to_string(img).strip()
        except Exception:
            result["ocr_text"] = None
            result["ocr_note"] = "OCR unavailable — tesseract not installed or image unreadable"
        return result
    except ImportError:
        raise HTTPException(status_code=500, detail="Pillow not installed")


# ─── Parse endpoints ──────────────────────────────────────────────────────────

@app.post("/api/documents/parse", tags=["documents"], summary="Parse an uploaded document")
async def parse_document_upload(file: UploadFile = File(...)):
    """Parse an uploaded file and return structured text/data."""
    content = await file.read()
    try:
        result = _parse_file_bytes(file.filename or "unknown", content)
        return {"success": True, "filename": file.filename, **result}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=str(e))


@app.post("/api/documents/parse-path", tags=["documents"], summary="Parse a document from a server path")
def parse_document_path(req: ParsePathRequest):
    """Parse a file from an absolute server-side path."""
    path = _safe_workspace_path(req.path)
    with open(path, "rb") as f:
        content = f.read()
    try:
        result = _parse_file_bytes(os.path.basename(path), content, req.encoding or "utf-8")
        return {"success": True, "path": path, **result}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=str(e))


@app.post("/api/documents/parse-image", tags=["documents"], summary="OCR an uploaded image")
async def parse_image_upload(file: UploadFile = File(...)):
    """Parse an uploaded image (OCR + metadata)."""
    content = await file.read()
    try:
        result = _parse_image_bytes(content, file.filename or "unknown")
        return {"success": True, "filename": file.filename, **result}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=str(e))


@app.post("/api/documents/parse-image-path", tags=["documents"], summary="OCR an image from a server path")
def parse_image_path_endpoint(req: ParseImagePathRequest):
    """Parse an image from an absolute server-side path (OCR + metadata)."""
    path = _safe_workspace_path(req.path)
    with open(path, "rb") as f:
        content = f.read()
    try:
        result = _parse_image_bytes(content, os.path.basename(path))
        return {"success": True, "path": path, **result}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=str(e))


@app.get("/api/documents/read", tags=["files"], summary="Read a text file from the workspace")
def read_text_file(
    path: str = Query(..., description="Absolute server path to a text file"),
    encoding: str = Query("utf-8"),
):
    """Read a plain-text file and return its content."""
    safe_path = _safe_workspace_path(path)
    try:
        with open(safe_path, "r", encoding=encoding, errors="replace") as f:
            content = f.read()
        return {
            "success": True,
            "path": safe_path,
            "content": content,
            "size": os.path.getsize(safe_path),
        }
    except Exception as e:
        raise HTTPException(status_code=422, detail=str(e))


@app.post("/api/documents/write", tags=["files"], summary="Write a text file to the workspace")
async def write_text_file(req: WriteFileRequest):
    """Write a plain-text file to the workspace or cloud if outside."""
    path, is_safe = _resolve_path(req.path, req.workspaceRoot)

    if not is_safe:
        if not req.userId:
            raise HTTPException(
                status_code=403,
                detail="Access denied: path is outside workspace and no userId provided for cloud upload",
            )
        enc = req.encoding or "utf-8"
        content_bytes = req.content.encode(enc)
        upload = await _upload_to_cloud(os.path.basename(path), content_bytes, req.userId)
        return {
            "success": True,
            "path": path,
            "url": upload["url"],
            "fileId": upload["fileId"],
            "storedAs": upload["storedAs"],
            "type": "cloud",
            "size": len(content_bytes),
        }

    try:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding=req.encoding or "utf-8") as f:
            f.write(req.content)
        return {
            "success": True,
            "path": path,
            "type": "local",
            "size": len(req.content),
        }
    except Exception as e:
        raise HTTPException(status_code=422, detail=str(e))


@app.post("/api/documents/delete", tags=["files"], summary="Delete a file from the workspace")
def delete_workspace_file(req: DeleteFileRequest):
    """Delete a file from the workspace."""
    safe_path = _safe_workspace_path(req.path, check_exists=True)
    try:
        if os.path.isdir(safe_path):
            import shutil
            shutil.rmtree(safe_path)
        else:
            os.remove(safe_path)
        return {"success": True, "path": safe_path}
    except Exception as e:
        raise HTTPException(status_code=422, detail=str(e))


from logger import get_nest_log_config, get_logger

logger = get_logger("Bootstrap")

if __name__ == "__main__":
    logger.info("🚀 Document Service is starting on: http://0.0.0.0:3009")
    uvicorn.run("main:app", host="0.0.0.0", port=3009, reload=False, log_config=get_nest_log_config())
